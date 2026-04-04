from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path

from dotenv import load_dotenv
from openai import OpenAI

import video_tranformer

ROOT_DIR = Path(__file__).resolve().parent
VIDEOS_DIR = ROOT_DIR / "videos"
JSON_DIR = ROOT_DIR / "json_data"
NOTES_DIR = ROOT_DIR / "output" / "notes"
FRAMES_DIR = NOTES_DIR / "frames"

DEFAULT_NOTES_MODEL = "openrouter/free"
SECTION_TARGET_SECONDS = 180.0
SECTION_MAX_SECONDS = 240.0
SECTION_MIN_CHARACTERS = 900
SECTION_MAX_CHARACTERS = 3200
OCR_SAMPLE_COUNT = 4
OCR_SIMILARITY_THRESHOLD = 0.9
MAX_SECTION_SLIDES = 2

load_dotenv()


@dataclass
class SlideCandidate:
    timestamp: float
    slide_text: str
    frame_path: Path | None


@dataclass
class NotesSection:
    index: int
    start: float
    end: float
    transcript: str
    slide_text: str
    slides: list[SlideCandidate]


def _normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _normalize_for_compare(text: str) -> str:
    text = _normalize_whitespace(text).lower()
    text = re.sub(r"[^a-z0-9 ]+", "", text)
    return text.strip()


def _looks_like_duplicate(candidate: str, existing_items: list[str], threshold: float = OCR_SIMILARITY_THRESHOLD) -> bool:
    normalized_candidate = _normalize_for_compare(candidate)
    if not normalized_candidate:
        return True

    for existing in existing_items:
        normalized_existing = _normalize_for_compare(existing)
        if not normalized_existing:
            continue
        if normalized_candidate == normalized_existing:
            return True
        if normalized_candidate in normalized_existing or normalized_existing in normalized_candidate:
            return True
        ratio = SequenceMatcher(None, normalized_candidate, normalized_existing).ratio()
        if ratio >= threshold:
            return True
    return False


def _dedupe_lines(lines: list[str], max_lines: int = 14) -> list[str]:
    cleaned: list[str] = []
    for line in lines:
        normalized = _normalize_whitespace(line)
        if len(normalized) < 3:
            continue
        if _looks_like_duplicate(normalized, cleaned):
            continue
        cleaned.append(normalized)
        if len(cleaned) >= max_lines:
            break
    return cleaned


def _dedupe_segments(segments: list[dict[str, object]]) -> list[dict[str, object]]:
    filtered: list[dict[str, object]] = []
    seen_recent: list[str] = []

    for segment in segments:
        text = _normalize_whitespace(str(segment.get("text", "")))
        if not text:
            continue
        if _looks_like_duplicate(text, seen_recent, threshold=0.94):
            continue
        filtered.append(
            {
                "start": float(segment["start"]),
                "end": float(segment["end"]),
                "text": text,
            }
        )
        seen_recent = (seen_recent + [text])[-8:]

    return filtered


def _build_sections(segments: list[dict[str, object]]) -> list[list[dict[str, object]]]:
    sections: list[list[dict[str, object]]] = []
    current: list[dict[str, object]] = []
    current_chars = 0
    section_start: float | None = None

    for segment in segments:
        start = float(segment["start"])
        end = float(segment["end"])
        text = str(segment["text"])

        if section_start is None:
            section_start = start

        proposed_chars = current_chars + len(text) + (1 if current else 0)
        proposed_duration = end - section_start

        should_split = False
        if current:
            if proposed_duration > SECTION_MAX_SECONDS:
                should_split = True
            elif current_chars >= SECTION_MIN_CHARACTERS and proposed_duration >= SECTION_TARGET_SECONDS:
                should_split = True
            elif proposed_chars > SECTION_MAX_CHARACTERS:
                should_split = True

        if should_split:
            sections.append(current)
            current = []
            current_chars = 0
            section_start = start

        current.append(segment)
        current_chars += len(text) + (1 if current_chars else 0)

    if current:
        sections.append(current)

    return sections


def _extract_ocr_text(video_path: Path, start: float, end: float, frame_output: Path | None = None) -> tuple[str, Path | None]:
    try:
        import cv2
        import pytesseract
    except ImportError:
        return "", None

    capture = cv2.VideoCapture(str(video_path))
    if not capture.isOpened():
        raise RuntimeError(f"Video is not readable for OCR: {video_path}")

    timestamps: list[float] = []
    if end <= start:
        timestamps = [start]
    else:
        step_count = max(1, OCR_SAMPLE_COUNT - 1)
        timestamps = [start + ((end - start) * idx / step_count) for idx in range(OCR_SAMPLE_COUNT)]

    best_frame = None
    unique_lines: list[str] = []

    try:
        for timestamp in timestamps:
            capture.set(cv2.CAP_PROP_POS_MSEC, max(0.0, timestamp) * 1000)
            ok, frame = capture.read()
            if not ok or frame is None:
                continue

            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            try:
                raw_text = pytesseract.image_to_string(gray)
            except Exception:
                raw_text = ""
            raw_lines = [re.sub(r"[^a-zA-Z0-9.,:;()/%+\- ]+", " ", line).strip() for line in raw_text.splitlines()]
            informative_lines = _dedupe_lines(raw_lines, max_lines=8)

            if len(informative_lines) > len(unique_lines):
                unique_lines = informative_lines
                best_frame = frame
    finally:
        capture.release()

    saved_frame_path: Path | None = None
    if best_frame is not None and frame_output is not None:
        frame_output.parent.mkdir(parents=True, exist_ok=True)
        cv2.imwrite(str(frame_output), best_frame)
        saved_frame_path = frame_output

    return "\n".join(unique_lines), saved_frame_path


def _extract_slide_candidates(video_path: Path, start: float, end: float, output_dir: Path, section_index: int) -> list[SlideCandidate]:
    try:
        import cv2
        import pytesseract
    except ImportError:
        return []

    capture = cv2.VideoCapture(str(video_path))
    if not capture.isOpened():
        raise RuntimeError(f"Video is not readable for OCR: {video_path}")

    if end <= start:
        timestamps = [start]
    else:
        step_count = max(1, OCR_SAMPLE_COUNT - 1)
        timestamps = [start + ((end - start) * idx / step_count) for idx in range(OCR_SAMPLE_COUNT)]

    raw_candidates: list[tuple[int, float, object, list[str]]] = []
    try:
        for timestamp in timestamps:
            capture.set(cv2.CAP_PROP_POS_MSEC, max(0.0, timestamp) * 1000)
            ok, frame = capture.read()
            if not ok or frame is None:
                continue

            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            try:
                raw_text = pytesseract.image_to_string(gray)
            except Exception:
                raw_text = ""

            raw_lines = [re.sub(r"[^a-zA-Z0-9.,:;()/%+\- ]+", " ", line).strip() for line in raw_text.splitlines()]
            informative_lines = _dedupe_lines(raw_lines, max_lines=10)
            if not informative_lines:
                continue

            score = sum(len(line) for line in informative_lines)
            raw_candidates.append((score, timestamp, frame, informative_lines))
    finally:
        capture.release()

    if not raw_candidates:
        return []

    raw_candidates.sort(key=lambda item: item[0], reverse=True)
    selected: list[SlideCandidate] = []
    seen_texts: list[str] = []
    output_dir.mkdir(parents=True, exist_ok=True)

    for _, timestamp, frame, informative_lines in raw_candidates:
        slide_text = "\n".join(informative_lines).strip()
        if _looks_like_duplicate(slide_text, seen_texts, threshold=0.88):
            continue

        frame_path = output_dir / f"section_{section_index:02d}_slide_{len(selected) + 1:02d}.jpg"
        cv2.imwrite(str(frame_path), frame)
        selected.append(
            SlideCandidate(
                timestamp=timestamp,
                slide_text=slide_text,
                frame_path=frame_path,
            )
        )
        seen_texts.append(slide_text)

        if len(selected) >= MAX_SECTION_SLIDES:
            break

    return selected


def _get_openrouter_client() -> tuple[OpenAI, str]:
    api_key = (os.getenv("OPENROUTER_API_KEY") or os.getenv("MIMO_API_KEY") or "").strip()
    if not api_key:
        raise EnvironmentError(
            "Missing API key. Add `OPENROUTER_API_KEY` or `MIMO_API_KEY` to `.env`."
        )

    model_name = (
        os.getenv("LECTURE_NOTES_MODEL")
        or os.getenv("OPENROUTER_MODEL")
        or DEFAULT_NOTES_MODEL
    ).strip()
    return OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key), model_name


def _build_notes_prompt(video_label: str, section: NotesSection) -> str:
    if section.slides:
        slide_blocks = []
        for idx, slide in enumerate(section.slides, start=1):
            slide_blocks.append(
                f"Slide {idx} at {_format_timestamp(slide.timestamp)}:\n{slide.slide_text}"
            )
        slide_block = "\n\n".join(slide_blocks)
    else:
        slide_block = "No reliable slide text was detected for this section."
    return f"""
You are creating high-quality lecture notes from a video section.

Video title:
{video_label}

Section timestamp:
{_format_timestamp(section.start)} to {_format_timestamp(section.end)}

Transcript evidence:
{section.transcript}

OCR slide evidence:
{slide_block}

Rules:
1. Use transcript evidence as the primary source of truth.
2. Use OCR only to support or clarify terminology, formulas, headings, or bullet points.
3. Remove repetition. If the speaker repeats a point, mention it once unless repetition changes meaning.
4. Do not invent missing details.
5. Prefer clean academic notes over chatty summaries.
6. Keep the notes specific, but avoid repeating the same sentence structure.
7. Return plain text only. Do not use markdown symbols like #, ##, **, *, -, or backticks.

Output structure exactly:
Main Topics:
...

Notes:
...

Key Takeaways:
...

Slide Terms / On-Screen Items:
...
""".strip()


def _call_llm(prompt: str) -> str:
    client, model_name = _get_openrouter_client()
    response = client.chat.completions.create(
        model=model_name,
        temperature=0.25,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.choices[0].message.content or "No notes returned."


def _format_timestamp(seconds: float) -> str:
    total_seconds = max(0, round(seconds))
    minutes, secs = divmod(total_seconds, 60)
    hours, minutes = divmod(minutes, 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def _transcript_path_for_video(relative_video_path: str) -> Path:
    stem = video_tranformer.video_output_stem(VIDEOS_DIR / relative_video_path, VIDEOS_DIR)
    return JSON_DIR / f"{stem}.json"


def _load_transcript_segments(relative_video_path: str) -> list[dict[str, object]]:
    transcript_path = _transcript_path_for_video(relative_video_path)
    if not transcript_path.exists():
        raise FileNotFoundError(
            f"Transcript not found for `{relative_video_path}`. Process that video before generating notes."
        )

    with transcript_path.open("r", encoding="utf-8") as file:
        payload = json.load(file)

    segments = payload.get("segments", [])
    if not segments:
        raise RuntimeError(f"No transcript segments found in {transcript_path.name}")
    return _dedupe_segments(segments)


def _output_stem(relative_video_path: str) -> str:
    rel_path = Path(relative_video_path)
    return video_tranformer.video_output_stem(VIDEOS_DIR / relative_video_path, VIDEOS_DIR) or rel_path.stem


def _video_label(relative_video_path: str) -> str:
    rel_path = Path(relative_video_path)
    title = re.sub(r"\s*\[[^\]]+\]\s*$", "", rel_path.stem).strip()
    if len(rel_path.parts) > 1:
        return f"{rel_path.parts[0]} / {title}"
    return title


def _strip_markdown_markers(text: str) -> str:
    cleaned = text.replace("**", "").replace("__", "").replace("`", "")
    cleaned = re.sub(r"^\s{0,3}#{1,6}\s*", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^\s*[-*+]\s+", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^\s*\d+\.\s+", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", cleaned)
    return cleaned.strip()


def _parse_note_sections(note_text: str) -> list[tuple[str, str]]:
    lines = [line.rstrip() for line in note_text.splitlines()]
    sections: list[tuple[str, list[str]]] = []
    current_title = "Notes"
    current_body: list[str] = []

    def flush() -> None:
        nonlocal current_body, current_title
        body = "\n".join(line for line in current_body if line.strip()).strip()
        if body:
            sections.append((current_title, body.split("\n")))
        current_body = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_body.append("")
            continue

        markdown_heading = re.match(r"^\s{0,3}#{1,6}\s*(.+?)\s*$", stripped)
        if markdown_heading:
            flush()
            current_title = _strip_markdown_markers(markdown_heading.group(1)) or "Notes"
            continue

        if stripped.endswith(":") and len(stripped) <= 40:
            flush()
            current_title = _strip_markdown_markers(stripped[:-1].strip()) or "Notes"
            continue

        current_body.append(_strip_markdown_markers(stripped))

    flush()
    return [(title, "\n".join(body).strip()) for title, body in sections]


def _write_docx(video_label: str, sections: list[NotesSection], notes: list[str], output_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return

    document = Document()
    document.add_heading(video_label, 1)

    for section, note_text in zip(sections, notes, strict=False):
        document.add_heading(
            f"Section {section.index} ({_format_timestamp(section.start)} - {_format_timestamp(section.end)})",
            2,
        )
        for slide_idx, slide in enumerate(section.slides, start=1):
            if slide.frame_path and slide.frame_path.exists():
                document.add_paragraph(
                    f"Important Slide {slide_idx} - {_format_timestamp(slide.timestamp)}"
                )
                document.add_picture(str(slide.frame_path), width=Inches(5.5))
            if slide.slide_text:
                document.add_paragraph(f"Slide text: {slide.slide_text}")

        for title, body in _parse_note_sections(note_text):
            document.add_heading(title, 3)
            for paragraph in [part.strip() for part in body.split("\n\n") if part.strip()]:
                document.add_paragraph(_strip_markdown_markers(paragraph))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _write_pdf(video_label: str, sections: list[NotesSection], notes: list[str], output_path: Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        return

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(str(output_path), pagesize=A4, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    heading_style = styles["Heading2"]
    subheading_style = styles["Heading3"]
    body_style = ParagraphStyle(
        "NotesBody",
        parent=styles["BodyText"],
        fontSize=10,
        leading=14,
        spaceAfter=8,
    )

    story = [Paragraph(video_label, title_style), Spacer(1, 0.2 * inch)]

    for idx, (section, note_text) in enumerate(zip(sections, notes, strict=False), start=1):
        story.append(
            Paragraph(
                f"Section {section.index} ({_format_timestamp(section.start)} - {_format_timestamp(section.end)})",
                heading_style,
            )
        )
        story.append(Spacer(1, 0.12 * inch))

        for slide_idx, slide in enumerate(section.slides, start=1):
            story.append(
                Paragraph(
                    f"Important Slide {slide_idx} - {_format_timestamp(slide.timestamp)}",
                    subheading_style,
                )
            )
            if slide.frame_path and slide.frame_path.exists():
                story.append(Image(str(slide.frame_path), width=5.5 * inch, height=3.2 * inch, kind="proportional"))
                story.append(Spacer(1, 0.08 * inch))
            if slide.slide_text:
                story.append(Paragraph(_strip_markdown_markers(slide.slide_text).replace("\n", "<br/>"), body_style))

        for title, body in _parse_note_sections(note_text):
            story.append(Paragraph(title, subheading_style))
            for paragraph in [part.strip() for part in body.split("\n\n") if part.strip()]:
                story.append(Paragraph(_strip_markdown_markers(paragraph).replace("\n", "<br/>"), body_style))

        if idx < len(sections):
            story.append(PageBreak())

    document.build(story)


def generate_notes_for_video(relative_video_path: str) -> dict[str, str | int]:
    video_path = VIDEOS_DIR / relative_video_path
    if not video_path.exists():
        raise FileNotFoundError(f"Video not found: {relative_video_path}")

    sections_data = _build_sections(_load_transcript_segments(relative_video_path))
    if not sections_data:
        raise RuntimeError(f"No note sections could be built for {relative_video_path}")

    output_stem = _output_stem(relative_video_path)
    video_label = _video_label(relative_video_path)
    section_notes: list[str] = []
    sections: list[NotesSection] = []

    for index, section_segments in enumerate(sections_data, start=1):
        start = float(section_segments[0]["start"])
        end = float(section_segments[-1]["end"])
        transcript = " ".join(str(segment["text"]) for segment in section_segments).strip()
        slides = _extract_slide_candidates(
            video_path,
            start,
            end,
            FRAMES_DIR / output_stem,
            index,
        )
        slide_text = "\n\n".join(slide.slide_text for slide in slides if slide.slide_text)

        section = NotesSection(
            index=index,
            start=start,
            end=end,
            transcript=transcript,
            slide_text=slide_text,
            slides=slides,
        )
        sections.append(section)
        section_notes.append(_call_llm(_build_notes_prompt(video_label, section)))

    markdown_output = NOTES_DIR / f"{output_stem}.md"
    markdown_lines = [f"# {video_label}", ""]
    for section, note_text in zip(sections, section_notes, strict=False):
        markdown_lines.extend(
            [
                f"## Section {section.index}",
                "",
                f"Timestamp: {_format_timestamp(section.start)} to {_format_timestamp(section.end)}",
                "",
            ]
        )
        if section.slides:
            markdown_lines.extend(
                [
                    "### Important Slides",
                    "",
                ]
            )
            for slide_idx, slide in enumerate(section.slides, start=1):
                markdown_lines.extend(
                    [
                        f"- Slide {slide_idx} at {_format_timestamp(slide.timestamp)}",
                        f"  - Text: {slide.slide_text or 'No OCR text available'}",
                        f"  - Image: {slide.frame_path}",
                    ]
                )
            markdown_lines.append("")
        markdown_lines.extend(
            [
                note_text.strip(),
                "",
            ]
        )
    markdown_output.parent.mkdir(parents=True, exist_ok=True)
    markdown_output.write_text("\n".join(markdown_lines).strip() + "\n", encoding="utf-8")

    docx_output = NOTES_DIR / f"{output_stem}.docx"
    _write_docx(video_label, sections, section_notes, docx_output)
    pdf_output = NOTES_DIR / f"{output_stem}.pdf"
    _write_pdf(video_label, sections, section_notes, pdf_output)

    return {
        "video": relative_video_path,
        "sections": len(sections),
        "markdown_path": str(markdown_output),
        "docx_path": str(docx_output) if docx_output.exists() else "",
        "pdf_path": str(pdf_output) if pdf_output.exists() else "",
    }


def generate_notes(video_files: list[str]) -> list[dict[str, str | int]]:
    if not video_files:
        raise ValueError("Provide at least one video to generate notes for.")
    return [generate_notes_for_video(video_file) for video_file in video_files]
